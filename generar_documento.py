import logging
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from config import RUTA_CV_DIR, RUTA_CV_ADAPTADO_WORD, RUTA_CV_ADAPTADO_PDF

logger = logging.getLogger("generar_documento")

DARK_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_BLUE = RGBColor(0x2E, 0x86, 0xDE)
MEDIUM_GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)


def _set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_colored_paragraph(doc, text, size=11, bold=False, color=None, alignment=None, space_after=4, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.08
    return p


def _add_section_header(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.space_before = Pt(0)
    run2 = p2.add_run("━" * 55)
    run2.font.name = "Calibri"
    run2.font.size = Pt(6)
    run2.font.color.rgb = ACCENT_BLUE


def _add_competency_chips(doc, competencies):
    if not competencies:
        return
    per_row = 2
    n = len(competencies)
    rows = (n + per_row - 1) // per_row
    table = doc.add_table(rows=rows, cols=per_row)
    table.autofit = True
    for i, comp in enumerate(competencies):
        row = i // per_row
        col = i % per_row
        cell = table.rows[row].cells[col]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(f"▸ {comp}")
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(1)
        _set_cell_shading(cell, "EDF2F9")

    _add_colored_paragraph(doc, "", size=2, space_after=2)


def _add_skill_badge(doc, label, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run_label = p.add_run(f"{label}: ")
    run_label.font.name = "Calibri"
    run_label.font.size = Pt(10)
    run_label.bold = True
    run_label.font.color.rgb = DARK_BLUE
    run_items = p.add_run(items)
    run_items.font.name = "Calibri"
    run_items.font.size = Pt(10)
    run_items.font.color.rgb = MEDIUM_GRAY


def _add_bullet(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    run = p.add_run("• ")
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = ACCENT_BLUE
    run.bold = True
    run2 = p.add_run(text)
    run2.font.name = "Calibri"
    run2.font.size = Pt(size)
    run2.font.color.rgb = MEDIUM_GRAY


def _add_project_entry(doc, project):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True
    run_title = p.add_run(project.get("title", ""))
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(11)
    run_title.bold = True
    run_title.font.color.rgb = DARK_BLUE
    tag = project.get("tag", "")
    if tag:
        run_tag = p.add_run(f"  [{tag}]")
        run_tag.font.name = "Calibri"
        run_tag.font.size = Pt(9)
        run_tag.font.color.rgb = ACCENT_BLUE
    gh = project.get("github", "")
    if gh:
        run_gh = p.add_run(f"  ·  {gh}")
        run_gh.font.name = "Calibri"
        run_gh.font.size = Pt(8)
        run_gh.font.color.rgb = LIGHT_GRAY

    subtitle = project.get("subtitle", "")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.space_before = Pt(0)
        run_sub = p2.add_run(subtitle)
        run_sub.font.name = "Calibri"
        run_sub.font.size = Pt(9)
        run_sub.font.italic = True
        run_sub.font.color.rgb = MEDIUM_GRAY

    for bullet in project.get("bullets", []):
        if bullet.strip():
            _add_bullet(doc, bullet.strip())

    impact = project.get("impact", "").strip()
    if impact:
        p_imp = doc.add_paragraph()
        p_imp.paragraph_format.space_after = Pt(1)
        p_imp.paragraph_format.space_before = Pt(1)
        p_imp.paragraph_format.left_indent = Cm(0.5)
        run_imp_label = p_imp.add_run("Impacto: ")
        run_imp_label.font.name = "Calibri"
        run_imp_label.font.size = Pt(9)
        run_imp_label.bold = True
        run_imp_label.font.color.rgb = ACCENT_BLUE
        run_imp_text = p_imp.add_run(impact)
        run_imp_text.font.name = "Calibri"
        run_imp_text.font.size = Pt(9)
        run_imp_text.font.color.rgb = MEDIUM_GRAY


def _add_simple_entry(doc, title, subtitle, year):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(3)
    run_title = p.add_run(title)
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(10.5)
    run_title.bold = True
    run_title.font.color.rgb = DARK_BLUE
    if subtitle:
        run_sub = p.add_run(f"  —  {subtitle}")
        run_sub.font.name = "Calibri"
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = MEDIUM_GRAY
    if year:
        run_year = p.add_run(f"  ({year})")
        run_year.font.name = "Calibri"
        run_year.font.size = Pt(9)
        run_year.font.color.rgb = LIGHT_GRAY


def _add_cert_entry(doc, name, org, year):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run("▸ ")
    r1.font.name = "Calibri"
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = ACCENT_BLUE
    r2 = p.add_run(name)
    r2.font.name = "Calibri"
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = DARK_BLUE
    if org:
        r3 = p.add_run(f"  —  {org}")
        r3.font.name = "Calibri"
        r3.font.size = Pt(9)
        r3.font.color.rgb = MEDIUM_GRAY
    if year:
        r4 = p.add_run(f"  ({year})")
        r4.font.name = "Calibri"
        r4.font.size = Pt(8.5)
        r4.font.color.rgb = LIGHT_GRAY


def _add_exp_entry(doc, exp):
    puesto = exp.get("puesto", "")
    empresa = exp.get("empresa", "")
    fechas = exp.get("fechas", "")
    label = puesto
    if empresa:
        label += f"  —  {empresa}"
    if fechas:
        label += f"  ({fechas})"
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(label)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.bold = True
    run.font.color.rgb = DARK_BLUE

    desc = exp.get("descripcion", "").strip()
    if desc:
        for linea in desc.split("\n"):
            linea = linea.strip()
            if linea:
                _add_bullet(doc, linea)


def _add_experience_section(doc, experiencias):
    if not experiencias:
        return
    _add_section_header(doc, "Experiencia Profesional")
    for exp in experiencias:
        _add_exp_entry(doc, exp)


def _slug_oferta(nombre_oferta):
    if not nombre_oferta:
        return ""
    slug = re.sub(r'[^\w\s-]', '', nombre_oferta.lower())
    slug = re.sub(r'[\s-]+', '_', slug).strip('_')
    return slug[:50]


def generar_word(cv, ruta=None, nombre_oferta=""):
    if ruta is None:
        if nombre_oferta:
            slug = _slug_oferta(nombre_oferta)
            ruta = RUTA_CV_DIR / f"CV_Adaptado_{slug}.docx"
        else:
            ruta = RUTA_CV_ADAPTADO_WORD
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.08

    # ─── HEADER ─────────────────────────────────────────────────────
    nombre = cv.get("nombre", "") or "Tu Nombre"
    titulo = cv.get("titulo", "Ingeniero en Telecomunicaciones | Ciberseguridad & IA")

    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True

    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]

    p_name = left_cell.paragraphs[0]
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_name = p_name.add_run(nombre)
    run_name.font.name = "Calibri"
    run_name.font.size = Pt(20)
    run_name.bold = True
    run_name.font.color.rgb = DARK_BLUE

    p_title = left_cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run(titulo)
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = ACCENT_BLUE

    _set_cell_shading(left_cell, "FFFFFF")
    _set_cell_shading(right_cell, "FFFFFF")

    contact_pairs = []
    email = cv.get("email", "")
    if email:
        contact_pairs.append(("📧", email))
    phone = cv.get("telefono", "")
    if phone:
        contact_pairs.append(("📞", phone))
    ciudad = cv.get("ciudad", "") or cv.get("location", "")
    if ciudad:
        contact_pairs.append(("📍", ciudad))
    linkedin = cv.get("linkedin", "")
    if linkedin:
        contact_pairs.append(("🔗", linkedin))
    github = cv.get("github", "")
    if github:
        contact_pairs.append(("💻", github))

    if not contact_pairs:
        contact_pairs.append(("📧", "email@example.com"))

    for icon, text in contact_pairs:
        p_contact = right_cell.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_contact.paragraph_format.space_after = Pt(0)
        p_contact.paragraph_format.space_before = Pt(1)
        r1 = p_contact.add_run(icon + " ")
        r1.font.name = "Calibri"
        r1.font.size = Pt(8)
        r2 = p_contact.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(8)
        r2.font.color.rgb = MEDIUM_GRAY

    right_cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(4)
    p_sep.paragraph_format.space_before = Pt(4)
    run_sep = p_sep.add_run("━" * 55)
    run_sep.font.name = "Calibri"
    run_sep.font.size = Pt(6)
    run_sep.font.color.rgb = ACCENT_BLUE

    # ─── RESUMEN ────────────────────────────────────────────────────
    resumen = cv.get("resumen", "").strip()
    if resumen:
        _add_section_header(doc, "Resumen Profesional")
        _add_colored_paragraph(doc, resumen, size=10, color=MEDIUM_GRAY, space_after=6)

    # ─── CORE COMPETENCIES ─────────────────────────────────────────
    comps = cv.get("core_competencies", [])
    if comps:
        _add_section_header(doc, "Competencias Clave")
        _add_competency_chips(doc, comps)

    # ─── TECHNICAL SKILLS ──────────────────────────────────────────
    skills = cv.get("skills", {})
    if skills:
        _add_section_header(doc, "Habilidades Técnicas")
        for cat, items in skills.items():
            if items:
                _add_skill_badge(doc, cat, items if isinstance(items, str) else ", ".join(str(x) for x in items))
        _add_colored_paragraph(doc, "", size=2, space_after=2)

    # ─── EXPERIENCE ────────────────────────────────────────────────
    experiencias = cv.get("experiencia", [])
    if experiencias:
        _add_experience_section(doc, experiencias)

    # ─── PROJECTS ──────────────────────────────────────────────────
    proyectos = cv.get("projects", [])
    if proyectos:
        _add_section_header(doc, "Proyectos Destacados")
        for proj in proyectos:
            _add_project_entry(doc, proj)

    # ─── EDUCATION ─────────────────────────────────────────────────
    education = cv.get("education", [])
    if education:
        _add_section_header(doc, "Formación Académica")
        for edu in education:
            if isinstance(edu, (list, tuple)):
                title = str(edu[0]) if len(edu) > 0 else ""
                sub = str(edu[1]) if len(edu) > 1 else ""
                yr = str(edu[2]) if len(edu) > 2 else ""
                _add_simple_entry(doc, title, sub, yr)

    # ─── CERTIFICATIONS ────────────────────────────────────────────
    certs = cv.get("certifications", [])
    if certs:
        _add_section_header(doc, "Certificaciones")
        for c in certs:
            if isinstance(c, (list, tuple)):
                name = str(c[0]) if len(c) > 0 else ""
                org = str(c[1]) if len(c) > 1 else ""
                yr = str(c[2]) if len(c) > 2 else ""
                _add_cert_entry(doc, name, org, yr)

    # ─── SOFT SKILLS ───────────────────────────────────────────────
    soft = cv.get("soft_skills", [])
    if soft:
        _add_section_header(doc, "Habilidades Blandas")
        soft_text = ", ".join(soft)
        _add_colored_paragraph(doc, soft_text, size=10, color=MEDIUM_GRAY, space_after=4)

    # ─── IDIOMAS ───────────────────────────────────────────────────
    idiomas = cv.get("idiomas", [])
    if idiomas:
        _add_section_header(doc, "Idiomas")
        _add_colored_paragraph(doc, ", ".join(idiomas), size=10, color=MEDIUM_GRAY, space_after=4)

    doc.save(str(ruta))
    logger.info("Word generated: %s", ruta)
    return ruta


def convertir_a_pdf(ruta_docx=None):
    if ruta_docx is None:
        ruta_docx = RUTA_CV_ADAPTADO_WORD
    ruta_docx = Path(ruta_docx)
    if not ruta_docx.exists():
        logger.error("Word file not found: %s", ruta_docx)
        return None

    ruta_pdf = ruta_docx.with_suffix(".pdf")

    try:
        result = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(ruta_docx.parent),
                str(ruta_docx),
            ],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0 and ruta_pdf.exists():
            logger.info("PDF generated: %s", ruta_pdf)
            return ruta_pdf
        logger.error("LibreOffice error: %s", result.stderr)
        return None
    except FileNotFoundError:
        logger.error("LibreOffice not installed")
        return None
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice timeout")
        return None
    except Exception as e:
        logger.error("PDF conversion error: %s", e)
        return None


def generar_pdf(cv, ruta_pdf=None, nombre_oferta=""):
    from generar_pdf_html import generar_pdf as _html_pdf
    if ruta_pdf is None:
        if nombre_oferta:
            slug = _slug_oferta(nombre_oferta)
            ruta_pdf = RUTA_CV_DIR / f"CV_Adaptado_{slug}.pdf"
        else:
            ruta_pdf = RUTA_CV_ADAPTADO_PDF
    return _html_pdf(cv, ruta_pdf=ruta_pdf)
